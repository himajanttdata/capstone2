from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from github_engine import GitHubEngine
from responder_engine import ResponderEngine
from docx import Document
from agent_graph import app as review_graph
from bs4 import BeautifulSoup
import markdown
import os
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import html

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

UPLOAD_FOLDER = "upload"
REPORT_FOLDER = "reports"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["REPORT_FOLDER"] = REPORT_FOLDER

@app.route('/')
def index():
  return send_from_directory(app.static_folder, 'index.html')


@app.route("/submit", methods=["POST"])
def submit():
    review_report = None

    if "repo_url" in request.form and request.form["repo_url"]:
        repo_url = request.form["repo_url"]
        github_engine = GitHubEngine()
        repo_name, all_chunks = github_engine.process_repository(repo_url)

    elif "file" in request.files:
        uploaded_file = request.files["file"]
        if uploaded_file.filename:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], uploaded_file.filename)
            uploaded_file.save(file_path)

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                file_content = f.read()

            from chunking_engine import Chunking
            chunking_engine = Chunking()
            all_chunks = chunking_engine.chunk_code(file_content)
            repo_name = uploaded_file.filename

    if all_chunks:
        result = review_graph.invoke({
            "repo": repo_name,
            "code": all_chunks
        })

        review_report = result["final_report"]

        # Save the review text
        report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(review_report)

    return jsonify({"success": True}) if review_report else jsonify({"error": "Failed to generate review"}), 200

@app.route("/analyze_selected", methods=["POST"])
def analyze_selected():
    data = request.get_json()
    file_urls = data.get("file_urls", [])
 
    if not file_urls:
        return jsonify({"success": False, "error": "No files provided."})
 
    file_contents_combined = ""
    for url in file_urls:
        response = requests.get(url)
        if response.ok:
            file_contents_combined += response.text + "\n\n"
        else:
            return jsonify({"success": False, "error": f"Failed to fetch {url}."})
 
    from chunking_engine import Chunking
    chunking_engine = Chunking()
    all_chunks = chunking_engine.chunk_code(file_contents_combined)
 
    if all_chunks:
        result = review_graph.invoke({
            "repo": "Selected_Files",
            "code": all_chunks
        })
 
        review_report = result["final_report"]
 
        # Save review
        report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(review_report)
 
        return jsonify({"success": True})
 
    return jsonify({"success": False, "error": "Failed to process selected files."})

# @app.route('/review2', methods=["GET"])
# def review2():
#     report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
#     if not os.path.exists(report_path):
#         return jsonify({"error": "No report found."}), 404

#     with open(report_path, "r", encoding="utf-8") as f:
#         markdown_content = f.read()

#     sections = []
#     current_section = None
#     for line in markdown_content.split('\n'):
#         if line.startswith('##'):
#             if current_section:
#                 sections.append(current_section)
#             current_section = {
#                 'title': line.replace('##', '').strip(),
#                 'body': ''
#             }
#         elif current_section:
#             current_section['body'] += line + '\n'

#     if current_section:
#         sections.append(current_section)

#     # Convert markdown body to HTML
#     for section in sections:
#         section['body'] = markdown.markdown(section['body'])

#     return jsonify(sections)



@app.route("/review2", methods=["GET"])
def review2():

    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    if not os.path.exists(report_path):
        return jsonify({"error": "No report found."}), 404

    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()

    soup = BeautifulSoup(review_text, "html.parser")
    sections = []
    current_section = None

    for tag in soup.find_all(["h3", "p", "ul", "ol", "pre", "blockquote"]):
        if tag.name == "h3":
            if current_section:
                sections.append(current_section)
            current_section = {
                "title": tag.get_text(strip=True),
                "body": ""
            }
        elif current_section:
            if tag.name == "pre":
                code_tag = tag.find("code")
                if code_tag:
                    raw_code = code_tag.decode_contents()
                    escaped_code = html.escape(raw_code)
                    lang_class = code_tag.get("class", [""])[0]
                    current_section["body"] += f"<pre><code class='{lang_class}'>{escaped_code}</code></pre>\n"
                else:
                    raw_code = tag.get_text()
                    current_section["body"] += f"<pre><code>{html.escape(raw_code)}</code></pre>\n"
            else:
                current_section["body"] += str(tag)

    if current_section:
        sections.append(current_section)

    return jsonify(sections)



        # Clean empty lines or lines starting with "-"
        # review_text = "\n".join(
        #     line for line in review_text.splitlines()
        #     if line.strip() != "" or not line.strip().startswith("-")
        # )

        # html = markdown.markdown(review_text, extensions=["extra"])
    #     soup = BeautifulSoup(review_text, "html.parser")

    #     sections = []
    #     current_section = None

    #     for tag in soup.find_all(["h2", "h3", "h4", "p", "ul", "ol", "pre", "blockquote"]):
    #         if tag.name == "h3":
    #             # Start a new section
    #             if current_section:
    #                 sections.append(current_section)

    #             current_title = tag.text.strip()
    #             current_section = {
    #                 "title": current_title,
    #                 "body": ""
    #             }
    #         elif current_section:
    #             current_section["body"] += f"{str(tag)}\n"

    #     # Append the last section
    #     if current_section:
    #         sections.append(current_section)
    #     # print("Sections: ",sections)

    # return jsonify(sections)



@app.route("/download_docx", methods=["GET"])
def download_docx():
    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    if not os.path.exists(report_path):
        return jsonify({"error": "Report not found"}), 404

    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()
    
    soup = BeautifulSoup(review_text, "html.parser")
    doc = Document()
    doc.add_heading("Code Review Report", level=1)

    for element in soup.find_all(True):
        tag_name = element.name.lower()
        text = element.get_text(strip=True)

        if not text:
            continue  # skip empty elements

        if tag_name == "h1":
            doc.add_heading(text, level=1)
        elif tag_name == "h2":
            doc.add_heading(text, level=2)
        elif tag_name == "h3":
            doc.add_heading(text, level=3)
        elif tag_name == "h4":
            doc.add_heading(text, level=4)
        elif tag_name == "p":
            doc.add_paragraph(text)
        elif tag_name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif tag_name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif tag_name == "pre":
            code_text = element.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
        elif tag_name == "code":
            run = doc.add_paragraph().add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    docx_path = os.path.join(REPORT_FOLDER, "VeriCODE_Review_Report.docx")
    doc.save(docx_path)

    return send_file(
        docx_path,
        as_attachment=True,
        download_name="VeriCODE_Review_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


    # review_html = markdown.markdown(review_text, extensions=["fenced_code", "codehilite"])
    # soup = BeautifulSoup(review_html, "html.parser")

    # doc = Document()
    # doc.add_heading("Code Review Report", level=1)

    # for element in soup.find_all(True):
    #     if element.name == "h1":
    #         doc.add_heading(element.get_text(), level=1)
    #     elif element.name == "h2":
    #         doc.add_heading(element.get_text(), level=2)
    #     elif element.name == "h3":
    #         doc.add_heading(element.get_text(), level=3)
    #     elif element.name == "p":
    #         doc.add_paragraph(element.get_text())
    #     elif element.name == "ul":
    #         for li in element.find_all("li"):
    #             doc.add_paragraph(li.get_text(), style="List Bullet")
    #     elif element.name == "ol":
    #         for li in element.find_all("li"):
    #             doc.add_paragraph(li.get_text(), style="List Number")
    #     elif element.name == "pre":
    #         code_text = element.get_text()
    #         para = doc.add_paragraph()
    #         run = para.add_run(code_text)
    #         run.font.name = "Courier New"
    #         run.font.size = Pt(10)
    #         para.paragraph_format.space_before = Pt(6)
    #         para.paragraph_format.space_after = Pt(6)

    #     elif element.name == "code":
    #         # p = doc.add_paragraph()
    #         # run = p.add_run(element.get_text())
    #         # run.font.name = 'Courier New'
    #         # run.font.size = Pt(10)
    #         run = doc.add_paragraph().add_run(element.get_text())
    #         run.font.name = "Courier New"
    #         run.font.size = Pt(10)

    # docx_path = os.path.join(REPORT_FOLDER, "VeriCODE_Review_Report.docx")
    # doc.save(docx_path)

    # return send_file(
    #     docx_path,
    #     as_attachment=True,
    #     download_name="VeriCODE_Review_Report.docx",  # ✅ Sets correct filename
    #     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # )



@app.route("/")
def health_check():
    return jsonify({"message": "Flask backend is running."})


# if __name__ == "__main__":
#     app.run(debug=True)

if __name__ == "__main__":
    # app.run(debug=True)
    port = int(os.environ.get('PORT',8000))
    app.run(host='0.0.0.0', port=port)

